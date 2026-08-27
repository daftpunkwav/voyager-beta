"""文档文本提取器:纯函数,按扩展名分发;worker 只做调度与事件。

设计约束:
- 提取失败抛 ExtractError(带用户可读原因),不静默返回空;
- 未知扩展名由 capabilities 层直接置 stored(存档语义),不进本模块;
- 章节切分只依赖文本本身:有结构(书签/标题样式)按结构,无结构按长度。
"""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree

#: 每章软上限(字符):无结构文档按此切块,超大章再按空行细分
_CHAPTER_TARGET = 8000


class ExtractError(Exception):
    """提取失败(加密/损坏/空文档),worker 捕获后置 failed。"""


@dataclass
class Section:
    section_no: int
    title: str
    page_start: int
    page_end: int
    text: str


def extract_sections(path: str | Path, ext: str) -> list[Section]:
    ext = ext.lower()
    if ext == ".pdf":
        return _from_pdf(Path(path))
    if ext == ".epub":
        return _from_epub(Path(path))
    if ext == ".docx":
        return _from_docx(Path(path))
    if ext in (".txt", ".md", ".markdown"):
        return _from_text(Path(path))
    raise ExtractError(f"不支持的解析格式: {ext}")


# ---------- PDF ----------

def _from_pdf(path: Path) -> list[Section]:
    import pypdfium2 as pdfium  # 延迟导入:非 PDF 路径不付出加载成本

    try:
        doc = pdfium.PdfDocument(str(path))
    except Exception as exc:  # 损坏/加密统一转用户可读错误
        raise ExtractError(f"PDF 打开失败: {exc}") from exc
    try:
        if doc.is_encrypted:
            raise ExtractError("PDF 已加密,暂不支持解析")
        page_texts: list[str] = []
        for page in doc:
            textpage = page.get_textpage()
            page_texts.append(textpage.get_text_bounded() or "")
            textpage.close()
            page.close()
        outline = _pdf_outline(doc)
        sections: list[Section] = []
        if outline:
            # 书签章:相邻书签页码之间为一章
            marks = sorted((page_no, title) for page_no, title in outline)
            for i, (page_no, title) in enumerate(marks):
                end = marks[i + 1][0] if i + 1 < len(marks) else len(page_texts)
                text = "\n".join(page_texts[page_no:end]).strip()
                if text:
                    sections.append(Section(len(sections) + 1, title,
                                            page_no + 1, end, text))
        if not sections:
            # 无书签:每 10 页一章(与阶段文档一致)
            for start in range(0, len(page_texts), 10):
                text = "\n".join(page_texts[start:start + 10]).strip()
                if text:
                    sections.append(Section(len(sections) + 1, "",
                                            start + 1,
                                            min(start + 10, len(page_texts)), text))
        if not sections:
            raise ExtractError("PDF 无可提取文本(可能是扫描件,OCR 暂不支持)")
        return sections
    finally:
        doc.close()


def _pdf_outline(doc) -> list[tuple[int, str]]:
    """书签 → (0 基页码, 标题);页码越界的书签丢弃。"""
    try:
        raw = doc.get_toc()
    except Exception:  # noqa: BLE001  # 书签解析失败按无书签处理
        return []
    marks: list[tuple[int, str]] = []
    count = len(doc)
    for item in raw:
        idx = item.page_index
        if idx is not None and 0 <= idx < count:
            marks.append((idx, item.title or ""))
    return marks


# ---------- EPUB ----------

_XHTML_NS = "{http://www.w3.org/1999/xhtml}"
_BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6", "li", "blockquote"}


def _from_epub(path: Path) -> list[Section]:
    try:
        zf = zipfile.ZipFile(path)
    except Exception as exc:
        raise ExtractError(f"EPUB 打开失败: {exc}") from exc
    with zf:
        spine_files = _epub_spine(zf)
        sections: list[Section] = []
        for i, name in enumerate(spine_files):
            try:
                root = ElementTree.fromstring(zf.read(name))
            except ElementTree.ParseError:
                continue
            title = _first_heading(root)
            text = _strip_xhtml(root)
            if not text.strip():
                continue
            sections.append(Section(len(sections) + 1, title, i + 1, i + 1, text))
        if not sections:
            raise ExtractError("EPUB 无可提取文本")
        return sections


def _epub_spine(zf: zipfile.ZipFile) -> list[str]:
    """OPF spine 顺序的 xhtml 文件路径列表;解析失败退化为全部 .x?html。"""
    names = zf.namelist()
    opf_name = next((n for n in names if n.endswith(".opf")), None)
    if opf_name is None:
        return [n for n in names if n.endswith((".xhtml", ".html", ".htm"))]
    base = opf_name.rsplit("/", 1)[0] + "/" if "/" in opf_name else ""
    try:
        root = ElementTree.fromstring(zf.read(opf_name))
    except ElementTree.ParseError:
        return [n for n in names if n.endswith((".xhtml", ".html", ".htm"))]
    manifest: dict[str, str] = {}
    for item in root.iter():
        if item.tag.endswith("}item") or item.tag == "item":
            iid, href = item.get("id"), item.get("href")
            if iid and href:
                manifest[iid] = href
    order: list[str] = []
    for item in root.iter():
        if item.tag.endswith("}itemref") or item.tag == "itemref":
            idref = item.get("idref")
            href = manifest.get(idref or "")
            if href:
                full = (base + href).lstrip("/")
                if full in names:
                    order.append(full)
    return order


def _first_heading(root: ElementTree.Element) -> str:
    for el in root.iter():
        if el.tag in {f"{_XHTML_NS}h{i}" for i in range(1, 7)} | {"h1", "h2", "h3",
                                                                 "h4", "h5", "h6"}:
            return (el.text or "").strip()[:100]
    return ""


def _strip_xhtml(root: ElementTree.Element) -> str:
    parts: list[str] = []

    def walk(el) -> None:
        tag = el.tag.rsplit("}", 1)[-1]
        if tag in ("script", "style"):
            return
        if el.text:
            parts.append(el.text)
        for child in el:
            walk(child)
            if child.tail:
                parts.append(child.tail)
        if tag in _BLOCK_TAGS:
            parts.append("\n")

    walk(root)
    return re.sub(r"[ \t]+", " ", "".join(parts)).strip()


# ---------- DOCX ----------

def _from_docx(path: Path) -> list[Section]:
    import docx  # 延迟导入

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise ExtractError(f"DOCX 打开失败: {exc}") from exc
    sections: list[Section] = []
    buf: list[str] = []
    title = ""

    def flush() -> None:
        nonlocal buf
        text = "\n".join(buf).strip()
        if text:
            sections.append(Section(len(sections) + 1, title, 0, 0, text))
        buf = []

    for para in document.paragraphs:
        style = (para.style.name or "") if para.style is not None else ""
        if style.startswith("Heading"):
            flush()
            title = para.text.strip()[:100]
            buf = [para.text]
        else:
            buf.append(para.text)
        if sum(len(b) for b in buf) > _CHAPTER_TARGET and title:
            flush()
            title = ""
    flush()
    if not sections:
        raise ExtractError("DOCX 无可提取文本")
    return sections


# ---------- 纯文本 ----------

def _from_text(path: Path) -> list[Section]:
    try:
        text = path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        raise ExtractError(f"文件读取失败: {exc}") from exc
    if not text.strip():
        raise ExtractError("文件为空")
    if len(text) <= _CHAPTER_TARGET:
        return [Section(1, "", 0, 0, text)]
    # 按空行切块,块聚到目标长度为一章(md 场景标题行可当章题)
    chunks = re.split(r"\n\s*\n", text)
    sections: list[Section] = []
    buf: list[str] = []
    size = 0
    for chunk in chunks:
        buf.append(chunk)
        size += len(chunk)
        if size >= _CHAPTER_TARGET:
            joined = "\n\n".join(buf).strip()
            sections.append(Section(len(sections) + 1,
                                    _md_title(joined), 0, 0, joined))
            buf, size = [], 0
    if buf:
        joined = "\n\n".join(buf).strip()
        if joined:
            sections.append(Section(len(sections) + 1, _md_title(joined), 0, 0, joined))
    return sections


def _md_title(text: str) -> str:
    """块首的 ATX 标题行作为章题(仅 md/txt 启发,取不到为空)。"""
    first = text.split("\n", 1)[0].strip()
    if first.startswith("#"):
        return first.lstrip("#").strip()[:100]
    return ""

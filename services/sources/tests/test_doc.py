"""doc 子模块测试:导入/解析管线(注入 parse_fn)/分章读取/校验/删除事件。

不触真实 PDF/docx 解析:worker 用 parse_fn 注入;提取器纯函数单测
(txt/epub/docx 构造与解析在本地完成,PDF 只测拒绝路径)。
"""

import asyncio
import zipfile
from pathlib import Path

import pytest
from platform_actor import ActorContext
from platform_capability import execute
from platform_contracts import LOCAL_USER, ActorKind, ActorRef, ServiceError
from platform_eventbus import EventBus, EventLog
from platform_secrets import SecretStore

from services.sources.capabilities import SourcesDeps, init_all, registry
from services.sources.modules.doc.extract import (
    ExtractError,
    _from_docx,
    _from_text,
    extract_sections,
)
from services.sources.modules.doc.store import DocStore
from services.sources.modules.doc.worker import DocWorker
from services.sources.modules.repo.store import RepoStore
from services.sources.modules.web.store import WebStore

USER_CTX = ActorContext(actor=LOCAL_USER)
AGENT_CTX = ActorContext(actor=ActorRef(kind=ActorKind.AGENT, id="agent.main", scopes=()))


@pytest.fixture()
def deps(tmp_path):
    log = EventLog(tmp_path / "events.db")
    d = SourcesDeps(
        repo_store=RepoStore(tmp_path / "repo.db"),
        doc_store=DocStore(tmp_path / "doc.db"),
        web_store=WebStore(tmp_path / "web.db"),
        secrets=SecretStore(tmp_path / "secrets.db", key_material="t"),
        bus=EventBus(log),
        repo_queue=asyncio.Queue(), doc_queue=asyncio.Queue(),
        workspace=tmp_path / "ws",
    )
    init_all(d)
    yield d, log
    d.repo_store.close()
    d.doc_store.close()
    d.web_store.close()
    log.close()


def _seed_file(ws: Path, name: str, content: bytes | str) -> Path:
    ws.mkdir(parents=True, exist_ok=True)
    p = ws / name
    if isinstance(content, str):
        p.write_text(content, encoding="utf-8")
    else:
        p.write_bytes(content)
    return p


class TestAddDocument:
    async def test_add_parseable_enqueues_and_emits_added(self, deps, tmp_path) -> None:
        d, log = deps
        src = _seed_file(tmp_path / "ws", "report.pdf", b"%PDF-1.4 fake")
        ref = await execute(registry, "add_document", USER_CTX,
                            {"file_path": str(src), "title": "季度报告",
                             "tags": ["工作"]})
        doc = d.doc_store.get(ref.job_id)
        assert doc["status"] == "parsing"  # 可解析格式入队等 worker
        assert doc["ext"] == ".pdf"
        assert d.doc_queue.qsize() == 1
        assert Path(doc["local_path"]).parent == Path(d.workspace) / "doc"
        types = [e.type for _, e in log.read_after()]
        assert "source.added" in types

    async def test_add_unknown_ext_stored_without_parse(self, deps, tmp_path) -> None:
        """未知扩展名=存档语义:收下入库,不入解析队列。"""
        d, _ = deps
        src = _seed_file(tmp_path / "ws", "dataset.zip", b"PK fake")
        ref = await execute(registry, "add_document", USER_CTX,
                            {"file_path": str(src)})
        doc = d.doc_store.get(ref.job_id)
        assert doc["status"] == "stored"
        assert d.doc_queue.qsize() == 0

    async def test_missing_file_and_outside_workspace(self, deps, tmp_path) -> None:
        with pytest.raises(ServiceError, match="文件不存在"):
            await execute(registry, "add_document", USER_CTX,
                          {"file_path": str(tmp_path / "ws" / "nope.pdf")})
        # jail 外无论是否存在都不透露
        with pytest.raises(ServiceError, match="workspace"):
            await execute(registry, "add_document", USER_CTX,
                          {"file_path": str(tmp_path / "nope.pdf")})
        outside = tmp_path / "outside"
        outside.mkdir()
        src = outside / "a.pdf"
        src.write_bytes(b"x")
        with pytest.raises(ServiceError, match="workspace"):
            await execute(registry, "add_document", USER_CTX,
                          {"file_path": str(src)})

    async def test_title_filename_sanitized(self, deps, tmp_path) -> None:
        """title 带路径分隔符/保留字符不得写逃逸 workspace/doc。"""
        d, _ = deps
        src = _seed_file(tmp_path / "ws", "a.md", b"# t")
        ref = await execute(registry, "add_document", USER_CTX,
                            {"file_path": str(src), "title": '../../evil "name"'})
        local = Path(d.doc_store.get(ref.job_id)["local_path"])
        assert local.parent == Path(d.workspace) / "doc"
        assert local.is_file()


class TestParseWorker:
    async def test_parse_then_ready_with_sections(self, deps, tmp_path) -> None:
        d, log = deps

        def fake_parse(path: Path, ext: str):
            return _from_text(Path(path))  # 复用真实文本切块

        worker = DocWorker(d.doc_store, EventBus(log), d.doc_queue,
                           tmp_path / "ws", parse_fn=fake_parse)
        await worker.start()
        src = _seed_file(tmp_path / "ws", "book.md", "# 章\n" + "内容" * 50)
        ref = await execute(registry, "add_document", USER_CTX,
                            {"file_path": str(src), "title": "手册"})  # 内部已入队
        for _ in range(50):
            await asyncio.sleep(0.02)
            if d.doc_store.get(ref.job_id)["status"] == "ready":
                break
        await worker.stop()
        doc = d.doc_store.get(ref.job_id)
        assert doc["status"] == "ready"
        outline = d.doc_store.sections_outline(ref.job_id)
        assert len(outline) >= 1
        section = d.doc_store.section(ref.job_id, outline[0]["section_no"])
        assert section["text"].startswith("# 章")
        types = [e.type for _, e in log.read_after(
            types=["source.ready", "task.progress"])]
        assert "source.ready" in types

    async def test_parse_failure_marks_failed(self, deps, tmp_path) -> None:
        d, log = deps

        def boom(path, ext):
            raise ExtractError("PDF 打开失败: 测试注入")

        worker = DocWorker(d.doc_store, EventBus(log), d.doc_queue,
                           tmp_path / "ws", parse_fn=boom)
        await worker.start()
        src = _seed_file(tmp_path / "ws", "bad.pdf", b"%PDF broken")
        ref = await execute(registry, "add_document", USER_CTX,
                            {"file_path": str(src)})  # 内部已入队
        for _ in range(50):
            await asyncio.sleep(0.02)
            if d.doc_store.get(ref.job_id)["status"] == "failed":
                break
        await worker.stop()
        doc = d.doc_store.get(ref.job_id)
        assert doc["status"] == "failed" and "PDF" in doc["error"]
        types = [e.type for _, e in log.read_after(types=["task.failed"])]
        assert types == ["task.failed"]

    async def test_remove_cleans_record_and_file(self, deps, tmp_path) -> None:
        d, log = deps

        def noop(path, ext):
            return []

        worker = DocWorker(d.doc_store, EventBus(log), d.doc_queue,
                           tmp_path / "ws", parse_fn=noop)
        await worker.start()
        src = _seed_file(tmp_path / "ws", "x.md", b"hi")
        ref = await execute(registry, "add_document", USER_CTX,
                            {"file_path": str(src)})
        local = d.doc_store.get(ref.job_id)["local_path"]
        await execute(registry, "remove_document", USER_CTX, {"doc_id": ref.job_id})
        await asyncio.sleep(0.05)  # worker 异步删文件
        await worker.stop()
        assert d.doc_store.get(ref.job_id) is None
        assert not Path(local).exists()
        types = [e.type for _, e in log.read_after(types=["source.removed"])]
        assert types == ["source.removed"]


class TestReadAndMeta:
    async def test_get_document_outline_and_section(self, deps, tmp_path) -> None:
        d, _ = deps
        src = _seed_file(tmp_path / "ws", "r.md", "正文".encode())
        ref = await execute(registry, "add_document", USER_CTX,
                            {"file_path": str(src), "title": "报告"})
        d.doc_store.replace_sections(ref.job_id, [
            {"section_no": 1, "title": "概述", "page_start": 1, "page_end": 2,
             "text": "第一章内容"},
            {"section_no": 2, "title": "方法", "page_start": 3, "page_end": 9,
             "text": "第二章内容"},
        ])
        d.doc_store.set_status(ref.job_id, "ready")
        detail = await execute(registry, "get_document", AGENT_CTX,
                               {"doc_id": ref.job_id})
        assert [s["title"] for s in detail["sections"]] == ["概述", "方法"]
        assert "text" not in detail["sections"][0]  # 大纲不含正文(§9.20)
        section = await execute(registry, "get_doc_section", AGENT_CTX,
                                {"doc_id": ref.job_id, "section_no": 2})
        assert section["text"] == "第二章内容"
        assert section["total_sections"] == 2
        with pytest.raises(ServiceError, match="章不存在"):
            await execute(registry, "get_doc_section", AGENT_CTX,
                          {"doc_id": ref.job_id, "section_no": 9})

    async def test_search_sections_with_snippet(self, deps, tmp_path) -> None:
        d, _ = deps
        src = _seed_file(tmp_path / "ws", "r.md", b"x")
        ref = await execute(registry, "add_document", USER_CTX,
                            {"file_path": str(src), "title": "AI 手册"})
        d.doc_store.replace_sections(ref.job_id, [
            {"section_no": 1, "title": "", "page_start": 0, "page_end": 0,
             "text": "前面铺垫。" * 20 + "注意力机制是核心。" + "后面内容。" * 20}])
        hits = await execute(registry, "search_documents", AGENT_CTX,
                             {"query": "注意力机制"})
        assert hits[0]["doc_id"] == ref.job_id and hits[0]["section_no"] == 1
        assert "注意力机制" in hits[0]["snippet"]

    async def test_set_meta_and_tag_validation(self, deps, tmp_path) -> None:
        d, _ = deps
        src = _seed_file(tmp_path / "ws", "m.md", b"x")
        ref = await execute(registry, "add_document", USER_CTX,
                            {"file_path": str(src), "title": "笔记集"})
        await execute(registry, "set_document_meta", USER_CTX,
                      {"doc_id": ref.job_id, "category": "课程", "tags": ["ml"],
                       "progress": "learning", "note": "在读"})
        doc = d.doc_store.get(ref.job_id)
        assert doc["category"] == "课程" and doc["tags"] == ["ml"]
        with pytest.raises(ServiceError, match="标签不合法"):
            await execute(registry, "set_document_meta", USER_CTX,
                          {"doc_id": ref.job_id, "tags": ['a"b']})

    async def test_list_filter_by_status_and_query(self, deps, tmp_path) -> None:
        d, _ = deps
        s1 = _seed_file(tmp_path / "ws", "a.md", b"x")
        s2 = _seed_file(tmp_path / "ws", "b.zip", b"z")
        r1 = await execute(registry, "add_document", USER_CTX,
                           {"file_path": str(s1), "title": "Alpha 手册"})
        await execute(registry, "add_document", USER_CTX,
                      {"file_path": str(s2), "title": "Beta 数据集"})
        d.doc_store.set_status(r1.job_id, "ready")
        ready_only = await execute(registry, "list_documents", USER_CTX,
                                   {"status": "ready"})
        assert [r["title"] for r in ready_only] == ["Alpha 手册"]
        by_query = await execute(registry, "list_documents", USER_CTX,
                                 {"query": "alpha"})
        assert len(by_query) == 1


class TestExtractor:
    def test_text_chunks_long_content(self, tmp_path) -> None:
        p = tmp_path / "long.md"
        p.write_text(("段落内容。\n\n" * 2000), encoding="utf-8")
        sections = extract_sections(p, ".md")
        assert len(sections) > 1  # 长文按空行切块分章
        assert [s.section_no for s in sections] == list(range(1, len(sections) + 1))

    def test_docx_heading_split(self, tmp_path) -> None:
        import docx
        document = docx.Document()
        document.add_paragraph("第一章", style="Heading 1")
        document.add_paragraph("第一章正文")
        document.add_paragraph("第二章", style="Heading 1")
        document.add_paragraph("第二章正文")
        p = tmp_path / "d.docx"
        document.save(str(p))
        sections = _from_docx(p)
        assert [s.title for s in sections] == ["第一章", "第二章"]

    def test_epub_spine_order(self, tmp_path) -> None:
        p = tmp_path / "b.epub"
        with zipfile.ZipFile(p, "w") as zf:
            zf.writestr("mimetype", "application/epub+zip")
            zf.writestr("content/ch1.xhtml",
                        "<html><body><h1>一</h1><p>甲</p></body></html>")
            zf.writestr("content/ch2.xhtml",
                        "<html><body><h1>二</h1><p>乙</p></body></html>")
            zf.writestr("content/book.opf",
                        '<package><manifest>'
                        '<item id="c1" href="ch1.xhtml"/>'
                        '<item id="c2" href="ch2.xhtml"/>'
                        '</manifest><spine>'
                        '<itemref idref="c1"/><itemref idref="c2"/>'
                        '</spine></package>')
        sections = extract_sections(p, ".epub")
        assert [s.title for s in sections] == ["一", "二"]

    def test_pdf_corrupt_raises_extract_error(self, tmp_path) -> None:
        p = tmp_path / "bad.pdf"
        p.write_bytes(b"not a pdf at all")
        with pytest.raises(ExtractError):
            extract_sections(p, ".pdf")

    def test_empty_text_raises(self, tmp_path) -> None:
        p = tmp_path / "e.txt"
        p.write_text("  \n  ", encoding="utf-8")
        with pytest.raises(ExtractError, match="为空"):
            extract_sections(p, ".txt")
